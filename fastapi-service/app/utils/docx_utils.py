import io
import base64
import logging
from typing import Optional, Dict, Any, List

logger = logging.getLogger(__name__)


# ─── TEXT EXTRACTION ─────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    """
    Extract all paragraph text from a DOCX file.
    Returns full text or None if extraction fails.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full_text = "\n".join(paragraphs)
        logger.info(f"DOCX extracted {len(full_text)} characters, {len(paragraphs)} paragraphs")
        return full_text if full_text else None
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        return None


def get_docx_metadata(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract DOCX metadata: author, title, headings, table count, image count.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        core = doc.core_properties

        headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]

        tables_count = len(doc.tables)

        # Count inline images (runs with picture elements)
        images_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                ):
                    images_count += 1

        return {
            "author":        core.author or "",
            "title":         core.title or "",
            "subject":       core.subject or "",
            "created":       str(core.created) if core.created else "",
            "modified":      str(core.modified) if core.modified else "",
            "headings":      headings,
            "tables_count":  tables_count,
            "images_count":  images_count,
            "paragraph_count": len(doc.paragraphs),
        }
    except Exception as e:
        logger.error(f"DOCX metadata extraction failed: {e}")
        return {}


def get_docx_headings(file_bytes: bytes) -> List[str]:
    """Extract all headings (H1, H2, H3) from the DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]
    except Exception as e:
        logger.error(f"Could not extract headings: {e}")
        return []


# ─── DOCX MANIPULATION ───────────────────────────────────────────────────────

def add_summary_to_docx(file_bytes: bytes, summary: str) -> Optional[bytes]:
    """
    Insert an AI-generated summary box at the TOP of the document,
    before any existing content. Returns modified DOCX bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(io.BytesIO(file_bytes))

        # Insert at position 0 (before all existing paragraphs)
        # We add 3 paragraphs at the top: title, summary, divider
        _insert_paragraph_at(doc, 0, "")                          # blank spacer after
        _insert_paragraph_at(doc, 0, summary, style_bold=False,
                              font_size=11, color=(50, 50, 50))   # summary text
        _insert_paragraph_at(doc, 0, "AI-Generated Summary",
                              style_bold=True, font_size=13,
                              color=(0, 70, 127))                  # heading

        output = io.BytesIO()
        doc.save(output)
        logger.info("Summary added to DOCX successfully")
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_summary_to_docx failed: {e}")
        return None


def add_tags_to_docx(file_bytes: bytes, tags: List[str]) -> Optional[bytes]:
    """
    Append AI tags as a section at the END of the document.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document(io.BytesIO(file_bytes))

        # Add at the end
        doc.add_paragraph("")  # spacer
        heading = doc.add_paragraph("AI-Generated Tags")
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 127)

        tags_para = doc.add_paragraph("  |  ".join(f"[{t}]" for t in tags))
        tags_para.runs[0].font.size = Pt(11)

        output = io.BytesIO()
        doc.save(output)
        logger.info("Tags added to DOCX successfully")
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_tags_to_docx failed: {e}")
        return None


def add_keywords_to_docx(file_bytes: bytes, keywords: List[str]) -> Optional[bytes]:
    """Append extracted keywords section at the END of the document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document(io.BytesIO(file_bytes))
        doc.add_paragraph("")
        h = doc.add_paragraph("AI-Extracted Keywords")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(13)
        h.runs[0].font.color.rgb = RGBColor(0, 70, 127)

        kw_para = doc.add_paragraph(", ".join(keywords))
        kw_para.runs[0].italic = True
        kw_para.runs[0].font.size = Pt(11)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_keywords_to_docx failed: {e}")
        return None


def replace_text_in_docx(file_bytes: bytes, find: str, replace: str) -> Optional[bytes]:
    """
    Find and replace text throughout the entire DOCX
    (paragraphs + table cells).
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        count = 0

        # Replace in paragraphs
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find in para.text:
                            for run in para.runs:
                                if find in run.text:
                                    run.text = run.text.replace(find, replace)
                                    count += 1

        output = io.BytesIO()
        doc.save(output)
        logger.info(f"replace_text: replaced '{find}' → '{replace}' in {count} runs")
        return output.getvalue()
    except Exception as e:
        logger.error(f"replace_text_in_docx failed: {e}")
        return None


def add_watermark_text_to_docx(file_bytes: bytes, watermark_text: str = "DRAFT") -> Optional[bytes]:
    """
    Add a text watermark to the header of each section in the DOCX.
    (True background watermarks require XML injection — this adds a
    visible header-based watermark which works universally.)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(io.BytesIO(file_bytes))

        for section in doc.sections:
            header = section.header
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.clear()
            run = para.add_run(f"[ {watermark_text} ]")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(180, 180, 180)  # light gray
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output = io.BytesIO()
        doc.save(output)
        logger.info(f"Watermark '{watermark_text}' added to DOCX")
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_watermark_text_to_docx failed: {e}")
        return None


def add_table_of_contents_to_docx(file_bytes: bytes) -> Optional[bytes]:
    """
    Insert a simple text-based Table of Contents (from headings)
    at the top of the DOCX. Word's auto-TOC requires macros;
    this inserts a clean static TOC from extracted headings.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document(io.BytesIO(file_bytes))

        headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]

        if not headings:
            logger.warning("No headings found for TOC generation")
            return file_bytes  # return original unchanged

        # Build TOC lines
        toc_lines = ["Table of Contents", ""] + [f"  • {h}" for h in headings] + [""]

        # Insert at top of document
        for i, line in enumerate(reversed(toc_lines)):
            bold = (line == "Table of Contents")
            color = (0, 70, 127) if bold else (50, 50, 50)
            _insert_paragraph_at(doc, 0, line, style_bold=bold,
                                  font_size=13 if bold else 11, color=color)

        output = io.BytesIO()
        doc.save(output)
        logger.info(f"TOC with {len(headings)} entries added to DOCX")
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_table_of_contents_to_docx failed: {e}")
        return None


def add_ai_analysis_report_to_docx(
    file_bytes: bytes,
    summary: str,
    tags: List[str],
    keywords: List[str],
    difficulty: str,
    quality_score: float,
    practice_questions: list
) -> Optional[bytes]:
    """
    Append a complete AI Analysis Report section at the END of the DOCX.
    This is the all-in-one function that adds everything in one shot.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(io.BytesIO(file_bytes))

        def add_heading(text, size=14, color=(0, 70, 127)):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(*color)

        def add_body(text, size=11):
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].font.size = Pt(size)

        # ── Divider ──
        doc.add_paragraph("─" * 60)
        add_heading("📊 AI Analysis Report", size=16, color=(0, 90, 160))
        doc.add_paragraph("")

        # ── Summary ──
        add_heading("Summary", size=13)
        add_body(summary)
        doc.add_paragraph("")

        # ── Difficulty & Quality ──
        add_heading("Difficulty & Quality", size=13)
        add_body(f"Difficulty Level : {difficulty}")
        add_body(f"Quality Score    : {round(quality_score * 100)}% content quality")
        doc.add_paragraph("")

        # ── Tags ──
        add_heading("Topic Tags", size=13)
        add_body("  |  ".join(f"[{t}]" for t in tags))
        doc.add_paragraph("")

        # ── Keywords ──
        add_heading("Key Terms", size=13)
        add_body(", ".join(keywords))
        doc.add_paragraph("")

        # ── Practice Questions ──
        if practice_questions:
            add_heading("Practice Questions", size=13)
            for i, q in enumerate(practice_questions, 1):
                q_text = q.question if hasattr(q, "question") else q.get("question", "")
                add_body(f"Q{i}. {q_text}", size=11)
                if hasattr(q, "options") and q.options:
                    for opt in q.options:
                        add_body(f"    {opt}", size=10)
                elif isinstance(q, dict) and q.get("options"):
                    for opt in q["options"]:
                        add_body(f"    {opt}", size=10)
            doc.add_paragraph("")

        output = io.BytesIO()
        doc.save(output)
        logger.info("Full AI Analysis Report added to DOCX")
        return output.getvalue()
    except Exception as e:
        logger.error(f"add_ai_analysis_report_to_docx failed: {e}")
        return None


# ─── HELPER ──────────────────────────────────────────────────────────────────

def _insert_paragraph_at(doc, index: int, text: str,
                          style_bold=False, font_size=11, color=(0, 0, 0)):
    """Insert a new paragraph at a specific position in the document body."""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    import copy

    # Create a new paragraph element
    new_para = doc.add_paragraph(text)
    if new_para.runs:
        run = new_para.runs[0]
        run.bold = style_bold
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color)

    # Move it to the desired index in the body XML
    body = doc.element.body
    # The newly added paragraph is at the end; move it to index
    para_elem = body[-1]  # just added paragraph
    body.remove(para_elem)
    body.insert(index, para_elem)


def bytes_to_base64(data: bytes) -> str:
    """Convert bytes to base64 string for JSON transport."""
    return base64.b64encode(data).decode("utf-8")
